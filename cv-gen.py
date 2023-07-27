from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import cv2

def create_cv():
    # Create a new Word document
    doc = Document()

    # Add your personal information to the CV
    full_name = "István Antal Löffler"
    date_of_birth = "2001.09.21\n"
    city_state_zip = "Budapest, 1143\n"
    email = "istvan.antal.loffler@gmail.com\n"
    phone = "+36307817072\n"
    info = date_of_birth+city_state_zip+email+phone


    doc.add_paragraph()
    doc.add_heading(full_name, level=1)
    doc.add_paragraph(info)

    # Add sections to the CV
    sections = {
        "Objective": "I am seeking an opportunity to apply my theoretical knowledge and technical skills in a dynamic and challenging environment. Eager to learn from experienced professionals and contribute to innovative projects that drive technological advancement.",
        "Education": """
        Budapest University of Technology and Economics        2021.09-ongoing
        1111 Budapest
        Faculty of Mechanical Engeneering
        BSc Mechatronics Engeering
        
        Szent Margit Gimnázium                                 2017.09-2021.06
        High School Diploma""",
        "Experience": "Currently I do not have any relevant work experience yet, but I am looking forward to learning and deepening my knowledge as an intern.",
        "Skills":
        """ My skills include:
            Languages: English C1, Hungarian
            Microsoft Office: MS office programs at a user level
            Technologies: C, C++, C#(beginner), Python, Matlab, NI Labview, FPWIN Pro7
            CAD: SolidWorks, SolidEdge, Fusion360(manufaturing)""",
    }

    for section, content in sections.items():
        doc.add_heading(section, level=2)
        doc.add_paragraph(content)

    # Save the CV as a Word document
    doc.save("generated_cv.docx")

create_cv()